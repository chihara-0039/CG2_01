from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "monthly_submission_docs"
ASSETS = OUT / "assets_current"
OUT.mkdir(exist_ok=True)
ACCOUNT = "LE3C_15_チハラ_シゴウ"
BLUE = RGBColor(31, 78, 121)
CYAN = RGBColor(25, 145, 170)
INK = RGBColor(28, 35, 43)
MUTED = RGBColor(90, 100, 112)

def font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = color

def setup(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(0.85)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    style.font.size = Pt(10.5)
    return doc

def rule(p, color="1F4E79"):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement("w:pBdr"); x = OxmlElement("w:bottom")
    x.set(qn("w:val"), "single"); x.set(qn("w:sz"), "10"); x.set(qn("w:color"), color)
    b.append(x); pPr.append(b)

def cover(doc, kicker, title, subtitle):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(58); r=p.add_run(kicker); font(r,12,True,CYAN)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); r=p.add_run(title); font(r,29,True,BLUE); rule(p)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(16); r=p.add_run(subtitle); font(r,14,False,MUTED)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(250); r=p.add_run("日本工学院専門学校　LE3C 15\nチハラ シゴウ\n2026年8月　月次提出"); font(r,11,False,MUTED)

def new_page(doc, title, lead=""):
    doc.add_page_break(); p=doc.add_paragraph(); r=p.add_run(title); font(r,22,True,BLUE); rule(p)
    if lead:
        p=doc.add_paragraph(); r=p.add_run(lead); font(r,11,False,MUTED)

def heading(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); font(r,14,True,CYAN)

def body(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.15
    r=p.add_run(text); font(r)

def bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3); r=p.add_run(item); font(r)

def picture(doc, name, width=6.7, caption=""):
    path=ASSETS/name
    if path.exists():
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(path), width=Inches(width))
        if caption:
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(caption); font(r,9,False,MUTED)

def footer(doc):
    for sec in doc.sections:
        p=sec.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        r=p.add_run(f"{ACCOUNT}  |  CG2_01"); font(r,8,False,MUTED)

def save(doc, suffix):
    footer(doc); path=OUT/f"{ACCOUNT}_{suffix}.docx"; doc.save(path); return path

def portfolio():
    d=setup(Document())
    cover(d,"PORTFOLIO 2026","自作エンジンで制作する\n3D探索アクション","移動・ジャンプ・ステージ攻略を軸に、制作ツールまで一体化した作品")
    new_page(d,"作品概要","プレイヤーを操作し、立体ステージを進んで星のゴールを目指す3D探索アクションです。")
    picture(d,"frame_05.png",6.8,"雷と照明が変化するステージ。画面上部の星がゴールの目印。")
    heading(d,"コアループ"); bullets(d,["タイトル画面からステージセレクトへ進み、遊ぶステージを選ぶ","ステージを観察し、移動とジャンプで足場を攻略する","カメラを回転・ズームして進行ルートを探す","星を取得し、専用演出とクリアシーンを経てステージセレクトへ戻る"])
    heading(d,"操作"); body(d,"キーボード／マウスとXboxコントローラーに対応。プレイヤー追従カメラと全体カメラを切り替えられます。")
    new_page(d,"ゲームを支える表現")
    picture(d,"frame_03.png",6.8,"Tempest Storm適用時。暗い空、雨、雷光、複数光源を組み合わせている。")
    bullets(d,["天候プリセット：通常、雨、雪、Tempest Storm","雨・雪の地面衝突エフェクト、雲、雷の大きさと長さのランダム化","天候に連動する天球色・明るさ・ライト設定","GPUパーティクル、ポストエフェクト、プレイヤー発光"])
    new_page(d,"制作ツールとワークフロー")
    body(d,"ゲーム本体だけでなく、ステージ制作・エフェクト調整・アニメーション確認を同じエンジン内で行えます。ImGuiの画面はUnityに近い構成へ整理しました。")
    bullets(d,["ステージエディタ：ブロック配置、保存、一覧表示、天候設定","Blender／外部JSONの読込と、ファイル更新時の再読込確認","エフェクトエディタと天候プリセット連携","スキニングエディタ：GPUスキニング、補間、骨表示、武器装着、左手GPUパーティクル"])
    picture(d,"frame_01.png",5.8,"ゲーム内に用意した操作説明オブジェクト。")
    new_page(d,"現在地と今後")
    heading(d,"現在できていること"); bullets(d,["BaseSceneを継承したタイトル／ゲーム／クリア各シーンとSceneManagerによる遷移","移動・ジャンプ・カメラ・落下復帰・スター取得判定","スター取得時の短い演出、COURSE CLEAR表示、継続する花火演出","外部ステージの読込、天候、照明、GPUパーティクル、ポストエフェクト"])
    heading(d,"次に強化すること"); bullets(d,["ステージ攻略のバリエーションとギミック追加","スター取得からクリアまでのカメラワークと演出調整","操作感、当たり判定、カメラ挙動の継続調整","企業向けデモとして短時間で魅力が伝わる導線作り"])
    body(d,"技術の数を増やすだけでなく、各機能をゲームの遊びへ結び付けることを今後の重点とします。")
    return save(d,"ポートフォリオ")

def program_doc():
    d=setup(Document())
    cover(d,"PROGRAM DOCUMENT","自作ゲームエンジン\nプログラム説明資料","3D探索アクションを題材にした、DirectX 12ベースのゲーム／ツール統合設計")
    new_page(d,"1. 全体設計","実行、編集、描画、ゲーム進行を役割別に分離し、MyGameはライフサイクル呼出しに集中させています。")
    bullets(d,["MyGame：Initialize / Update / Draw / Finalizeの窓口","GameRuntime：ゲーム全体の更新と描画の統括","各Controller：ステージ、天候、カメラ、UI、シーン遷移を担当","Manager群：テクスチャ、モデル、パーティクル、ライト等のリソース管理"])
    heading(d,"狙い"); body(d,"巨大化していたMyGame.cppの処理を分割し、機能の所在を明確化。追加機能がゲーム本体へ波及しにくい構成にしました。")
    new_page(d,"2. プレイヤーとカメラ")
    bullets(d,["移動・ジャンプ・落下時のスポーン／中継地点復帰","追従カメラと全体カメラの切替","マウスとXboxコントローラーによる回転・ズーム","スターとの接触判定、取得演出、クリアシーンへの進行"])
    body(d,"入力はキーボードとXInputを共通のプレイヤー更新へ渡し、操作デバイスによってゲームロジックが分岐しすぎないようにしています。")
    picture(d,"frame_05.png",6.5,"ゲームプレイ画面。ステージ、ゴール、天候演出を同時に確認できる。")
    new_page(d,"3. ステージ制作と外部連携")
    bullets(d,["ImGuiステージエディタで配置・削除・回転・保存","保存済みステージと外部レベルの一覧表示","Blenderから出力したJSONをゲームとエディタの両方で読込","起動中のファイル変更を検知し、確認後に再読込"])
    body(d,"外部ツールで編集した結果を短い反復で確認できるため、ステージ制作の試行回数を増やせます。")
    new_page(d,"4. 天候・照明・エフェクト")
    bullets(d,["天候プリセットで空色、明るさ、ライト、パーティクルを一括変更","雨・雪の衝突地点で専用パーティクルを生成","Tempest Stormの雲・風雨・雷をゲームシーンへ統合","複数光源によりプレイヤーライトと雷光を同時使用","シーン遷移時に天候とポストエフェクトをNormalへ復帰"])
    picture(d,"frame_03.png",6.5,"暗天候時のステージ。環境光と雷演出がゲーム空間へ反映される。")
    new_page(d,"5. GPUスキニングとアニメーション")
    bullets(d,["Compute Shaderで頂点スキニングを実行","複数メッシュ／複数マテリアルのモデルを扱う","アニメーション間を補間し、急なTポーズへの復帰を抑制","ボーン・ジョイント・名前・軸をデバッグ表示","右手ジョイントへ武器、左手ジョイントへGPUパーティクルを追従"])
    body(d,"CPUで全頂点を毎フレーム変形せず、GPU側に処理を寄せることで、モデル切替や複数描画へ拡張しやすい構成を目指しています。")
    new_page(d,"6. シーン管理・描画・安定性")
    bullets(d,["BaseSceneを継承したTitleScene／GameClearSceneを実装","SceneManagerとSceneFactoryで生成・更新・描画・終了処理を統一","タイトル、ステージセレクト、ゲーム、クリアの遷移を明示","クリア文字の表示完了後、画面奥で花火を次シーンまで継続","GPUパーティクル、ポストエフェクト、複数ライトの状態をシーン単位で管理"])
    heading(d,"今後の改善"); body(d,"GPU負荷の計測、モデル切替の回帰テスト、ステージデータ検証、ゲーム側の完了定義に沿った自動チェックを追加し、作品としての安定性を高めます。")
    return save(d,"プログラム説明資料")

if __name__ == "__main__":
    print(portfolio())
    print(program_doc())
